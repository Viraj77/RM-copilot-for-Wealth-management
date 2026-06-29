import streamlit as st
import time
import pandas as pd
from client_db import get_all_clients, get_client_profile
from agent import create_agent
from langgraph.checkpoint.memory import MemorySaver
from tools import market_data_tool, rag_retriever, suitability_checker
from analytics import evaluate_retrieval, evaluate_suitability_gate, evaluate_faithfulness, compare_single_vs_multihop
from ingest import ingest_all, ingest_file, extract_metadata

# Update and initialize the Chroma DB on startup exactly once if it does not exist
@st.cache_resource
def startup_ingest():
    import os
    if not os.path.exists("chroma_db") or len(os.listdir("chroma_db")) == 0:
        st.write("Initializing knowledge base index...")
        ingest_all()
        return "ingested"
    else:
        st.write("Preserving existing Chroma DB. Skipping startup ingestion.")
        return "preserved"

# Force a re-check if chroma_db is missing (clears stale cache from previous session)
import os as _os
if not _os.path.exists("chroma_db") or len(_os.listdir("chroma_db")) == 0:
    startup_ingest.clear()

startup_ingest()

def render_client_brief(state):
    brief = state.get("final_brief")
    if not brief:
        st.warning("No brief available to render.")
        return
        
    st.subheader("📄 Client Investment Brief")
    
    # Compliance Status Banner
    status = brief.get("compliance_status", "Cleared")
    status_style = "gate-cleared" if status == "Cleared" else ("gate-review" if status == "Needs Review" else "gate-blocked")
    
    st.markdown(f"""
    <div class='compliance-gate {status_style}'>
        Compliance Status: {status}
    </div>
    """, unsafe_allow_html=True)
    
    # Metadata
    col1, col2 = st.columns(2)
    with col1:
        st.markdown(f"**Client ID**: {brief.get('client_id')}")
    with col2:
        st.markdown(f"**Risk Profile**: `{brief.get('risk_profile')}`")
        
    st.markdown("### Portfolio Summary")
    st.write(brief.get("portfolio_summary"))
    
    st.markdown("### Recommendations")
    recos = brief.get("recommendations", [])
    if recos:
        for i, reco in enumerate(recos):
            st.markdown(f"#### 💡 Recommendation {i+1}: {reco.get('idea')}")
            st.markdown(f"**Rationale**: {reco.get('rationale')}")
            st.markdown(f"**Suitability Assessment**: {reco.get('suitability')}")
            if reco.get("citations"):
                st.markdown(f"*Citations*: {', '.join(reco.get('citations'))}")
            st.write("")
    else:
        st.info("No recommendations provided in this brief.")
        
    st.markdown("### RM Talking Points")
    tps = brief.get("talking_points", [])
    if tps:
        for tp in tps:
            st.markdown(f"- {tp}")
    else:
        st.info("No talking points generated.")
        
    st.markdown("---")
    st.caption(brief.get("disclaimer", "Decision support for RMs, not automated advice."))


def render_freeform_response(state):
    st.subheader("💬 Analytical Response")
    st.markdown(
        "<span style='background:#7C3AED;color:white;padding:3px 10px;"
        "border-radius:12px;font-size:12px;font-weight:600;'>🔍 Response Mode: Analytical (Free-form)</span>",
        unsafe_allow_html=True
    )
    
    st.markdown(state["free_form_response"])
    
    if state.get("retrieved_evidence"):
        with st.expander("📚 Cited Sources & Retrieved Evidence", expanded=False):
            for i, doc in enumerate(state["retrieved_evidence"]):
                st.markdown(f"**[{i+1}] Source: {doc['doc_id']}** (Category: {doc['type']}, Sensitivity: {doc['sensitivity']})")
                st.info(doc["content"])
                
    st.markdown("---")
    st.caption("This is an analytical summary for RM reference. Not client-specific advice.")

# Set page configuration with premium title and icon
st.set_page_config(
    page_title="Horizon Relationship Manager Copilot",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Premium Style Tweaks
st.markdown("""
<style>
    .reportview-container {
        background: #f0f2f6;
    }
    .main-header {
        font-family: 'Outfit', 'Inter', sans-serif;
        color: #1E3A8A;
        font-weight: 700;
        margin-bottom: 0px;
    }
    .sub-header {
        color: #4B5563;
        font-weight: 400;
        margin-top: 0px;
    }
    .compliance-gate {
        padding: 15px;
        border-radius: 8px;
        font-weight: bold;
        margin-bottom: 15px;
    }
    .gate-cleared {
        background-color: #D1FAE5;
        color: #065F46;
        border: 1px solid #34D399;
    }
    .gate-review {
        background-color: #FEF3C7;
        color: #92400E;
        border: 1px solid #FBBF24;
    }
    .gate-blocked {
        background-color: #FEE2E2;
        color: #991B1B;
        border: 1px solid #FCA5A5;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("<h1 class='main-header'>Horizon Wealth Management</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-header'>Relationship Manager Copilot & Compliance Advisor</p>", unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["💼 Relationship Manager Workspace", "📊 Analytics & Benchmarks (Part B)", "📚 Knowledge Base Manager"])

with tab1:
    # Sidebar: Client selection & Parameters
    st.sidebar.header("Client & RM Selection")
    clients = get_all_clients()
    client_options = {f"{c['client_id']} - {c['name']}": c['client_id'] for c in clients}
    selected_label = st.sidebar.selectbox("Select Target Client", list(client_options.keys()))
    client_id = client_options[selected_label]

    client_profile = get_client_profile(client_id)

    # Display RM info
    st.sidebar.markdown("### Assigned Relationship Manager")
    st.sidebar.text(f"RM Name: {client_profile['rm_name']}")
    st.sidebar.text(f"Research Entitlement: Tier {client_profile['rm_research_tier']}")

    # Default simulation parameters to False/None/0.0 as simulation functionality is removed
    simulate_trade = False
    product_code = None
    allocation_amount = 0.0

    # Display Client Profile & Portfolio
    col1, col2 = st.columns([1, 2])

    with col1:
        st.subheader("Client Information")
        st.markdown(f"**Client Name**: {client_profile['name']}")
        st.markdown(f"**Client ID**: {client_profile['client_id']}")
        st.markdown(f"**Risk Profile**: `{client_profile['risk_profile']}`")
        
        st.subheader("Current Holdings")
        total_val = 0.0
        for h in client_profile["holdings"]:
            st.markdown(f"- **{h['product_name']}** ({h['product_code']})")
            st.markdown(f"  Allocation: **USD {h['allocation_amount']:,.2f}** | Class: {h['asset_class']}")
            total_val += h['allocation_amount']
        st.markdown(f"**Total Portfolio Value**: **USD {total_val:,.2f}**")

    with col2:
        st.subheader("Copilot Workspace")
        
        # Pre-defined query templates or custom input
        query_templates = [
            f"Prepare talking points for client {client_id}'s quarterly review",
            f"Summarize the portfolio risk for client {client_id}",
            f"Is SCN-US-24 suitable for client {client_id}?",
            f"What is the house research view on Technology and Applied AI sectors?"
        ]
        
        query_selection = st.selectbox("Choose a request template", ["Custom Request"] + query_templates)
        
        if query_selection == "Custom Request":
            query_text = st.text_area("Enter your custom request for the Copilot:", value=f"Prepare talking points for client {client_id}'s quarterly review.")
        else:
            query_text = query_selection
            
        run_btn = st.button("Execute Copilot Agent", type="primary")

    # Initialize Session State for LangGraph execution
    if "checkpointer" not in st.session_state:
        st.session_state.checkpointer = MemorySaver()
    if "graph_state" not in st.session_state:
        st.session_state.graph_state = None
    if "thread_id" not in st.session_state:
        st.session_state.thread_id = None
    if "agent_run" not in st.session_state:
        st.session_state.agent_run = False
    if "interrupt_state" not in st.session_state:
        st.session_state.interrupt_state = False
    if "response_mode" not in st.session_state:
        st.session_state.response_mode = "structured"

    # Handler to run the agent
    if run_btn:
        st.session_state.agent_run = True
        st.session_state.interrupt_state = False
        
        # Unique thread ID for checkpointing
        st.session_state.thread_id = f"thread_{client_id}_{int(time.time())}"
        
        agent = create_agent(st.session_state.checkpointer)
        
        # force_structured=True when the trade simulation sidebar is active
        # (product_code + allocation_amount are set) — compliance gate must always run
        force_structured = simulate_trade and bool(product_code)

        # Initial State
        initial_state = {
            "query": query_text,
            "client_id": client_id,
            "product_code": product_code,
            "allocation_amount": allocation_amount,
            "force_structured": force_structured,
            "response_mode": "structured",   # default, overwritten by classify_intent_node
            "client_profile": client_profile,
            "retrieved_evidence": [],
            "compliance_status": "Cleared",
            "escalated": False,
            "review_notes": None,
            "final_brief": None,
            "free_form_response": None,
        }
        
        config = {"configurable": {"thread_id": st.session_state.thread_id}}
        
        # Adapt spinner message for custom requests
        spinner_msg = (
            "🔍 Classifying query and routing to best response mode..."
            if query_selection == "Custom Request"
            else "Executing agentic workflow..."
        )

        with st.spinner(spinner_msg):
            current_state = agent.invoke(initial_state, config)
            st.session_state.graph_state = current_state
            st.session_state.response_mode = current_state.get("response_mode", "structured")
            
            next_steps = agent.get_state(config).next
            if "human_review" in next_steps:
                st.session_state.interrupt_state = True
                st.warning("⚠️ COMPLIANCE GATE ESCALATION: The recommendation requires Relationship Manager/Compliance review.")
            else:
                mode_label = "Analytical (Free-form)" if st.session_state.response_mode == "freeform" else "Structured Brief"
                st.success(f"✅ Workflow completed successfully. Response mode: **{mode_label}**")

    # If Agent is currently running or paused
    if st.session_state.agent_run and st.session_state.graph_state:
        state = st.session_state.graph_state
        agent = create_agent(st.session_state.checkpointer)
        config = {"configurable": {"thread_id": st.session_state.thread_id}}
        
        with col2:
            # 1. Display intermediate steps/logs
            with st.expander("🔍 View Execution Steps", expanded=True):
                # Show response mode badge
                mode = state.get("response_mode", "structured")
                mode_color = "#2563EB" if mode == "structured" else "#7C3AED"
                mode_label = "Structured Brief" if mode == "structured" else "Analytical (Free-form)"
                st.markdown(
                    f"<span style='background:{mode_color};color:white;padding:3px 10px;"
                    f"border-radius:12px;font-size:12px;font-weight:600;'>Mode: {mode_label}</span>",
                    unsafe_allow_html=True
                )
                st.write("**Plan**: ", state.get("plan", "No planning notes."))
                
                if state.get("suitability_report"):
                    report = state["suitability_report"]
                    status_style = "gate-cleared" if report["status"] == "Cleared" else ("gate-review" if report["status"] == "Needs Review" else "gate-blocked")
                    st.markdown(f"""
                    <div class='compliance-gate {status_style}'>
                        Compliance Status: {report['status']}<br/>
                        Violations: {', '.join(report.get('violations', [])) or 'None'}<br/>
                        Reasons: {', '.join(report.get('reasons', [])) or 'None'}
                    </div>
                    """, unsafe_allow_html=True)
                    
                if state.get("retrieved_evidence"):
                    st.markdown("### Retrieved Knowledge (RAG)")
                    for i, doc in enumerate(state["retrieved_evidence"]):
                        st.markdown(f"**[{i+1}] Source: {doc['doc_id']}** (Category: {doc['type']}, Sensitivity: {doc['sensitivity']})")
                        st.info(doc["content"][:300] + "...")

            # 2. If paused at human review, show override / approval panel
            if st.session_state.interrupt_state:
                st.subheader("🧑‍💼 Compliance Oversight & Override Panel")
                st.error(f"Review required. Current programmatic status is: **{state.get('compliance_status')}**")
                notes = st.text_area("Compliance Review Notes / Rationale for Override", placeholder="Enter justification...")
                
                col_approve, col_reject = st.columns(2)
                
                with col_approve:
                    if st.button("Acknowledge & Force Approve Override", type="primary"):
                        agent.update_state(config, {"review_notes": notes, "compliance_status": "Cleared"}, as_node="human_review")
                        with st.spinner("Finalizing brief after override..."):
                            final_state = agent.invoke(None, config)
                            st.session_state.graph_state = final_state
                            st.session_state.interrupt_state = False
                            st.session_state.response_mode = "structured"  # override always produces structured brief
                            st.success("Override completed. Brief synthesized.")
                            st.rerun()
                            
                with col_reject:
                    if st.button("Uphold Block & Terminate"):
                        st.session_state.interrupt_state = False
                        st.error("Recommendation rejected by RM. Workflow terminated.")
                        st.session_state.agent_run = False

            # 3. Render output — branch on response_mode
            if not st.session_state.interrupt_state:
                current_mode = state.get("response_mode", "structured")

                if current_mode == "freeform" and state.get("free_form_response"):
                    render_freeform_response(state)
                elif state.get("final_brief"):
                    render_client_brief(state)

with tab2:
    st.subheader("Analytical Evaluation & Benchmarking (Part B)")
    st.markdown("This tab displays evaluation results on the Golden Set for retrieval accuracy, suitability gate metrics, recommendation groundedness, and single-shot vs multi-hop retrieval comparisons.")
    
    # Custom Evaluation Dataset File Uploader
    uploaded_eval_file = st.file_uploader(
        "Import Custom Evaluation Dataset (JSON)",
        type=["json"],
        help="Upload a custom eval_dataset.json file to run the tests against it."
    )
    
    eval_data = None
    if uploaded_eval_file is not None:
        try:
            import json
            eval_data = json.load(uploaded_eval_file)
            st.success("Successfully loaded custom evaluation dataset!")
        except Exception as e:
            st.error(f"Failed to parse custom JSON file: {e}")
            
    col_a, col_b = st.columns(2)
    with col_a:
        evaluate_retrieval(eval_data)
    with col_b:
        evaluate_suitability_gate(eval_data)
        
    st.divider()
    
    col_c, col_d = st.columns(2)
    with col_c:
        evaluate_faithfulness(eval_data)
    with col_d:
        compare_single_vs_multihop()

with tab3:
    st.subheader("📚 Knowledge Base Manager")
    st.markdown("Browse current documents and ingest new resources into the Chroma DB Vector Store.")
    
    import os
    docs_dir = "docs"
    
    st.write("### 📂 Current Knowledge Source Documents")
    
    file_list = []
    if os.path.exists(docs_dir):
        for root, dirs, files in os.walk(docs_dir):
            for file in files:
                file_path = os.path.join(root, file)
                # Skip temp or system files if any
                if file.startswith(".") or "__" in file_path:
                    continue
                    
                try:
                    size_kb = os.path.getsize(file_path) / 1024
                    
                    # Load text sample depending on extension to extract metadata
                    sample_text = ""
                    ext = os.path.splitext(file)[1].lower()
                    if ext == ".txt":
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            sample_text = f.read(1000)
                    elif ext == ".csv":
                        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                            sample_text = f.read(1000)
                    elif ext == ".docx":
                        import docx
                        doc = docx.Document(file_path)
                        sample_text = "\n".join([p.text for p in doc.paragraphs[:5]])[:1000]
                    elif ext == ".pdf":
                        import pypdf
                        reader = pypdf.PdfReader(file_path)
                        if reader.pages:
                            sample_text = reader.pages[0].extract_text()[:1000]
                            
                    meta = extract_metadata(sample_text, file_path)
                    
                    file_list.append({
                        "Filename": file,
                        "Doc ID": meta.get("doc_id", "N/A"),
                        "Type": meta.get("type", "Unknown"),
                        "Sensitivity": meta.get("sensitivity", "Public"),
                        "Source": meta.get("source", "Unknown"),
                        "Date": meta.get("date", "N/A"),
                        "Size": f"{size_kb:.1f} KB",
                        "Path": file_path
                    })
                except Exception as e:
                    file_list.append({
                        "Filename": file,
                        "Doc ID": "N/A",
                        "Type": "Unknown",
                        "Sensitivity": "Public",
                        "Source": "Unknown",
                        "Date": "N/A",
                        "Size": f"{os.path.getsize(file_path)/1024:.1f} KB",
                        "Path": file_path
                    })
                
    if file_list:
        df_docs = pd.DataFrame(file_list)
        st.dataframe(df_docs[["Filename", "Doc ID", "Type", "Sensitivity", "Source", "Date", "Size"]], use_container_width=True)
        
        # Detail view
        selected_file = st.selectbox("Select a file to inspect", [f["Path"] for f in file_list])
        if selected_file:
            with st.expander("🔍 View Raw Sample Content", expanded=False):
                ext = os.path.splitext(selected_file)[1].lower()
                try:
                    if ext == ".txt" or ext == ".csv":
                        with open(selected_file, "r", encoding="utf-8", errors="ignore") as f:
                            st.text(f.read(3000))
                    elif ext == ".docx":
                        import docx
                        doc = docx.Document(selected_file)
                        st.text("\n".join([p.text for p in doc.paragraphs[:20]]))
                    elif ext == ".pdf":
                        import pypdf
                        reader = pypdf.PdfReader(selected_file)
                        text_pages = [p.extract_text() or "" for p in reader.pages[:3]]
                        st.text("\n--- Page Break ---\n".join(text_pages))
                except Exception as e:
                    st.error(f"Could not read preview: {e}")
    else:
        st.info("No documents found in knowledge base directory.")
        
    st.divider()
    
    st.write("### ➕ Upload & Ingest New Document")
    col1_uploader, col2_selector = st.columns(2)
    with col1_uploader:
        uploaded_file = st.file_uploader(
            "Choose a document file to add to knowledge base",
            type=["txt", "pdf", "docx", "csv"]
        )
    with col2_selector:
        category = st.selectbox(
            "Document Category / Subfolder",
            ["Product Guides", "Compliance", "Research", "General Root"]
        )
        
    if uploaded_file is not None:
        folder_mapping = {
            "Product Guides": "docs/product_guides",
            "Compliance": "docs/compliance",
            "Research": "docs/research",
            "General Root": "docs"
        }
        target_folder = folder_mapping[category]
        os.makedirs(target_folder, exist_ok=True)
        target_path = os.path.join(target_folder, uploaded_file.name)
        
        st.info(f"Target Save Location: `{target_path}`")
        
        if st.button("Index & Persist to Chroma DB", type="primary"):
            try:
                # Save the file first
                with open(target_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                st.success(f"File saved to `{target_path}` successfully!")
                
                # Now run the ingest function
                with st.spinner("Parsing, chunking, and indexing new document..."):
                    res = ingest_file(target_path)
                    
                st.balloons()
                st.success(f"✅ Ingestion successful! Document split into {res['chunks_count']} chunks and persisted to Chroma DB.")
                
                # Show parsed metadata
                st.write("**Extracted Metadata Details**:")
                st.json(res["metadata"])
                
                # Rerun to update the file list
                time.sleep(2)
                st.rerun()
            except Exception as e:
                st.error(f"Ingestion failed: {e}")
