import os
import re
import csv
import hashlib
import pandas as pd
import docx
import pypdf
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Load environment variables
load_dotenv()

PERSIST_DIR = "chroma_db"
COLLECTION_NAME = "wealth_mgmt_knowledge"

def get_embeddings():
    return OpenAIEmbeddings(model="text-embedding-3-small")

def clean_text(text):
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def extract_metadata(text, file_path):
    """
    Dynamically extracts document metadata (doc_id, type, date, source, sensitivity)
    from text content or file path, without assuming a fixed structure.
    """
    meta = {}
    
    # 1. Extract doc_id
    doc_id = None
    # Check text for Document ID or Doc ID
    doc_id_match = re.search(r'(?:Document ID|Doc ID|DocID|Doc_ID)\s*:\s*([A-Za-z0-9\-]+)', text, re.IGNORECASE)
    if doc_id_match:
        doc_id = doc_id_match.group(1).strip()
    else:
        # Check filename for standard pattern like PG-001 or CMP-002
        fn = os.path.basename(file_path)
        fn_match = re.search(r'([A-Za-z0-9]+-[0-9]+)', fn)
        if fn_match:
            doc_id = fn_match.group(1).strip()
        else:
            # Fallback: assign a unique ID based on MD5 hash of path
            path_hash = hashlib.md5(file_path.encode('utf-8')).hexdigest()[:8]
            doc_id = f"DOC-{path_hash.upper()}"
    meta['doc_id'] = doc_id

    # 2. Extract type from subfolder or text
    doc_type = "document"
    path_lower = file_path.lower()
    if "product" in path_lower:
        doc_type = "product"
    elif "compliance" in path_lower:
        doc_type = "compliance"
    elif "research" in path_lower:
        doc_type = "research"
    else:
        # Search text for clues
        type_match = re.search(r'(?:Type|Category)\s*:\s*([A-Za-z]+)', text, re.IGNORECASE)
        if type_match:
            doc_type = type_match.group(1).strip().lower()
    meta['type'] = doc_type

    # 3. Extract date
    date_val = "N/A"
    date_match = re.search(r'(?:Effective Date|Publication Date|Date)\s*:\s*([\d]{4}-[\d]{2}-[\d]{2})', text, re.IGNORECASE)
    if date_match:
        date_val = date_match.group(1).strip()
    else:
        # Check any YYYY-MM-DD in text
        any_date = re.search(r'(\b\d{4}-\d{2}-\d{2}\b)', text)
        if any_date:
            date_val = any_date.group(1).strip()
    meta['date'] = date_val

    # 4. Extract source
    source_val = "Unknown Source"
    source_match = re.search(r'(?:Author|Source|Publisher)\s*:\s*([^\n]+)', text, re.IGNORECASE)
    if source_match:
        source_val = source_match.group(1).strip()
    else:
        # Set default based on type
        if doc_type == "product":
            source_val = "Product Governance"
        elif doc_type == "compliance":
            source_val = "Horizon Compliance Office"
        elif doc_type == "research":
            source_val = "Horizon Wealth Research Desk"
    meta['source'] = source_val

    # 5. Extract sensitivity
    sens_val = "Public"
    sens_match = re.search(r'(?:Sensitivity|Access Level)\s*:\s*([^\n]+)', text, re.IGNORECASE)
    if sens_match:
        sens_text = sens_match.group(1).lower()
        if "restricted" in sens_text or "internal" in sens_text or "tier 2" in sens_text or "tier 3" in sens_text:
            sens_val = "Restricted"
    else:
        # General scan of the first 1000 characters for restricted keywords
        sample = text[:1000].lower()
        if "restricted" in sample or "internal" in sample or "confidential" in sample or "tier 2" in sample:
            sens_val = "Restricted"
    meta['sensitivity'] = sens_val
    
    return meta

def load_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return clean_text(text)

def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            full_text.append(row_text)
            
    return clean_text("\n".join(full_text))

def load_pdf(file_path):
    reader = pypdf.PdfReader(file_path)
    full_text = []
    for page in reader.pages:
        full_text.append(page.extract_text() or "")
    return clean_text("\n\n".join(full_text))

def load_csv_rows(file_path):
    """
    Parses CSV and returns a list of row descriptions, along with their metadata.
    """
    df = pd.read_csv(file_path)
    rows_docs = []
    
    for idx, row in df.iterrows():
        row_dict = row.to_dict()
        
        # We can form a text block of all columns for ingestion
        content_items = []
        for k, v in row_dict.items():
            if pd.notna(v):
                content_items.append(f"{k}: {v}")
        row_text = "\n".join(content_items)
        
        # Extract metadata from this row text + file path
        meta = extract_metadata(row_text[:1000], file_path)
        
        # Remove metadata columns from page content to keep clean if present
        cleaned_content_items = []
        for k, v in row_dict.items():
            if k not in ['doc_id', 'type', 'date', 'source', 'sensitivity'] and pd.notna(v):
                cleaned_content_items.append(f"{k}: {v}")
        content_str = f"Document: {meta['doc_id']} ({meta['type'].upper()})\n" + "\n".join(cleaned_content_items)
        
        rows_docs.append((content_str, meta))
        
    return rows_docs

def ingest_all():
    docs_dir = "docs"
    
    if not os.path.exists(docs_dir):
        print(f"Error: Docs directory '{docs_dir}' not found.")
        return

    # Clear existing Chroma DB to prevent duplicate chunks on re-ingestion / app restart
    import shutil
    if os.path.exists(PERSIST_DIR):
        try:
            shutil.rmtree(PERSIST_DIR)
            print(f"Cleared existing Chroma DB at {PERSIST_DIR}")
        except Exception as e:
            print(f"Error clearing {PERSIST_DIR}: {e}")
        
    all_documents = []
    
    # Dynamic folder scanner
    for root, dirs, files in os.walk(docs_dir):
        for file in files:
            file_path = os.path.join(root, file)
            ext = os.path.splitext(file)[1].lower()
            
            print(f"Scanning and loading {file_path}...")
            
            try:
                if ext == ".csv":
                    # CSV returns multiple row documents
                    rows = load_csv_rows(file_path)
                    for content, meta in rows:
                        all_documents.append(Document(page_content=content, metadata=meta))
                else:
                    # Txt, docx, pdf return a single text block
                    if ext == ".txt":
                        text = load_txt(file_path)
                    elif ext == ".docx":
                        text = load_docx(file_path)
                    elif ext == ".pdf":
                        text = load_pdf(file_path)
                    else:
                        print(f"Skipping unsupported file type: {file}")
                        continue
                        
                    meta = extract_metadata(text, file_path)
                    all_documents.append(Document(page_content=text, metadata=meta))
                    
            except Exception as e:
                print(f"Error parsing file {file_path}: {e}")
                
    print(f"Total raw document objects loaded: {len(all_documents)}")
    
    # Chunking: Size not exceeding 500 tokens, 50 token overlap.
    # We use RecursiveCharacterTextSplitter.from_tiktoken_encoder to split by tokens.
    token_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=500,
        chunk_overlap=50,
        encoding_name="cl100k_base" # standard tiktoken encoding for GPT-4/o-mini
    )
    
    final_chunks = token_splitter.split_documents(all_documents)
    print(f"Split raw documents into {len(final_chunks)} token-bounded chunks.")
    
    # Print sample metadata and chunks for verification
    if final_chunks:
        print("Sample Chunk Metadata:")
        print(final_chunks[0].metadata)
        print("Sample Chunk Text (first 200 chars):")
        print(final_chunks[0].page_content[:200] + "...")
        
    # Store in Chroma
    embeddings = get_embeddings()
    print("Embedding chunks and storing in Chroma...")
    db = Chroma.from_documents(
        final_chunks, 
        embeddings, 
        persist_directory=PERSIST_DIR,
        collection_name=COLLECTION_NAME
    )
    db.persist()
    print(f"Successfully persisted ChromaDB to {PERSIST_DIR}")

def ingest_file(file_path):
    """
    Ingests, chunks, extracts metadata, embeds, and adds a single document file to the existing Chroma DB.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    all_documents = []
    
    if ext == ".csv":
        rows = load_csv_rows(file_path)
        for content, meta in rows:
            all_documents.append(Document(page_content=content, metadata=meta))
    else:
        if ext == ".txt":
            text = load_txt(file_path)
        elif ext == ".docx":
            text = load_docx(file_path)
        elif ext == ".pdf":
            text = load_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
            
        meta = extract_metadata(text, file_path)
        all_documents.append(Document(page_content=text, metadata=meta))
        
    token_splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=500,
        chunk_overlap=50,
        encoding_name="cl100k_base"
    )
    final_chunks = token_splitter.split_documents(all_documents)
    
    embeddings = get_embeddings()
    db = Chroma(
        persist_directory=PERSIST_DIR,
        embedding_function=embeddings,
        collection_name=COLLECTION_NAME
    )
    db.add_documents(final_chunks)
    db.persist()
    print(f"Successfully added {len(final_chunks)} chunks from {file_path} to ChromaDB.")
    
    return {
        "chunks_count": len(final_chunks),
        "metadata": all_documents[0].metadata if all_documents else {}
    }

if __name__ == "__main__":
    ingest_all()
